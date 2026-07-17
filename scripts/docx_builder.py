"""
Clone-based DOCX builder.

Each output paragraph is a deep-copied paragraph from a reference template,
with its runs cleared and replaced with new text. This preserves the exact
pPr (paragraph properties) and the rPr (run properties) from the reference,
producing output that looks identical.

Template indices are auto-detected from each template document rather than
hard-coded, so the builder works correctly with reference_resume_comp.docx,
referece_resume_exp.docx, and the original reference_resume.docx.
"""

import re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE     = ROOT / "reference_resume.docx"
TEMPLATE_COMP = ROOT / "reference_resume_comp.docx"
TEMPLATE_EXP  = ROOT / "referece_resume_exp.docx"   # note: typo in filename is intentional

# ── Template loading & index detection ───────────────────────────────────────

_tmpl_cache = {}      # path_str → {name: paragraph_object}
_active_template = None   # set to str path for the duration of each build call


def _detect_indices(paras):
    """Auto-detect paragraph indices for a template document."""
    indices = {'name': 0, 'contact': 1, 'heading': 2, 'bullet': 3}

    # company / role / subsection: first three Normal paragraphs after
    # the PROFESSIONAL EXPERIENCE Heading 1.
    prof_idx = next(
        (i for i, p in enumerate(paras)
         if p.style.name == 'Heading 1' and 'PROFESSIONAL' in p.text),
        None,
    )
    if prof_idx is not None:
        normals = [i for i in range(prof_idx + 1, len(paras))
                   if paras[i].style.name == 'Normal']
        if len(normals) >= 1: indices['company']    = normals[0]
        if len(normals) >= 2: indices['role']       = normals[1]
        if len(normals) >= 3: indices['subsection'] = normals[2]

    # body / conf: first non-list Body Text paragraph
    for i, p in enumerate(paras):
        if p.style.name == 'Body Text':
            ppr = p._p.find(qn('w:pPr'))
            numpr = ppr.find(qn('w:numPr')) if ppr is not None else None
            if numpr is None:
                indices.setdefault('body', i)
                indices.setdefault('conf', i)
                break

    # Fall-backs for old reference_resume.docx
    indices.setdefault('company', 9)
    indices.setdefault('role',    10)
    indices.setdefault('body',    21)
    indices.setdefault('conf',    41)
    return indices


def _load_template(path_str):
    """Load and cache paragraph objects keyed by style name."""
    if path_str not in _tmpl_cache:
        doc = Document(path_str)
        paras = doc.paragraphs
        idx = _detect_indices(paras)
        _tmpl_cache[path_str] = {
            k: paras[v] for k, v in idx.items() if v < len(paras)
        }
    return _tmpl_cache[path_str]


def _tpara(name):
    tp = _active_template or str(TEMPLATE)
    return _load_template(tp).get(name)


# ── Run property helpers ──────────────────────────────────────────────────────

def _rpr(tmpl_name, run_idx=0):
    """Return a fresh deep-copy of a template run's rPr element."""
    para = _tpara(tmpl_name)
    if para is None:
        para = _tpara('role') or _tpara('body')
    if para is None:
        return OxmlElement('w:rPr')
    runs = para.runs
    if not runs or run_idx >= len(runs):
        return OxmlElement('w:rPr')
    rpr = runs[run_idx]._r.find(qn('w:rPr'))
    return deepcopy(rpr) if rpr is not None else OxmlElement('w:rPr')


def _set_bold(rpr, bold):
    b = rpr.find(qn('w:b'))
    if bold and b is None:
        b = OxmlElement('w:b')
        rpr.insert(0, b)
    elif not bold and b is not None:
        rpr.remove(b)


def _set_italic(rpr, italic):
    i = rpr.find(qn('w:i'))
    if italic and i is None:
        i = OxmlElement('w:i')
        rpr.append(i)
    elif not italic and i is not None:
        rpr.remove(i)


def _make_run(text, rpr_elem, bold=None, italic=None, underline=False):
    """Build a <w:r> element using a cloned rPr."""
    rpr = deepcopy(rpr_elem)
    if bold is not None:
        _set_bold(rpr, bold)
    if italic is not None:
        _set_italic(rpr, italic)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rpr.append(u)
    r = OxmlElement('w:r')
    r.append(rpr)
    t = OxmlElement('w:t')
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    return r


def _make_hyperlink(doc_part, text, url, rpr_elem):
    """Build a <w:hyperlink> element."""
    r_id = doc_part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    rpr = deepcopy(rpr_elem)
    rs = OxmlElement('w:rStyle')
    rs.set(qn('w:val'), 'Hyperlink')
    rpr.insert(0, rs)
    r = OxmlElement('w:r')
    r.append(rpr)
    t = OxmlElement('w:t')
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    hl.append(r)
    return hl


# ── Inline markdown parser ────────────────────────────────────────────────────

_INLINE = re.compile(
    r'\*\*\*(.+?)\*\*\*'
    r'|\*\*(.+?)\*\*'
    r'|\*(.+?)\*'
    r'|\[([^\]]+)\]\{\.underline\}'
    r'|\[([^\]]+)\]\(([^\)]*)\)'
)


def _parse_inline(text):
    """Return list of (text, bold, italic, underline, href)."""
    tokens, last = [], 0
    for m in _INLINE.finditer(text):
        if m.start() > last:
            tokens.append((text[last:m.start()], False, False, False, None))
        if m.group(1):
            tokens.append((m.group(1), True, True, False, None))
        elif m.group(2):
            tokens.append((m.group(2), True, False, False, None))
        elif m.group(3):
            tokens.append((m.group(3), False, True, False, None))
        elif m.group(4):
            tokens.append((m.group(4), False, False, True, None))
        elif m.group(5):
            tokens.append((m.group(5), False, False, False, m.group(6)))
        last = m.end()
    if last < len(text):
        tokens.append((text[last:], False, False, False, None))
    return tokens


# ── Core paragraph insertion ──────────────────────────────────────────────────

def _clone_pPr(tmpl_name):
    """Return a deep-copy of the template paragraph's pPr."""
    para = _tpara(tmpl_name)
    if para is None:
        return None
    ppr = para._p.find(qn('w:pPr'))
    return deepcopy(ppr) if ppr is not None else None


def _new_para_xml(tmpl_name):
    """Create a fresh <w:p> with the template's pPr but no runs."""
    p = OxmlElement('w:p')
    ppr = _clone_pPr(tmpl_name)
    if ppr is not None:
        p.append(ppr)
    return p


def _add_inline_runs(p_xml, text, base_rpr, doc_part,
                     force_bold=False, force_italic=False):
    """Parse inline markdown in text and append runs to p_xml."""
    base_has_bold   = base_rpr.find(qn('w:b')) is not None
    base_has_italic = base_rpr.find(qn('w:i')) is not None

    for seg, b, i, u, href in _parse_inline(text):
        if not seg:
            continue
        want_bold   = force_bold  or b
        want_italic = force_italic or i

        bold_arg   = True if want_bold   else (False if base_has_bold   else None)
        italic_arg = True if want_italic else (False if base_has_italic else None)

        if href:
            p_xml.append(_make_hyperlink(doc_part, seg, href, base_rpr))
        else:
            p_xml.append(_make_run(seg, base_rpr,
                                   bold=bold_arg, italic=italic_arg, underline=u))


def _insert(doc, tmpl_name, text, doc_part=None,
            force_bold=False, force_italic=False):
    """Insert a new paragraph (cloned from template) with the given text."""
    if doc_part is None:
        doc_part = doc.part
    p_xml    = _new_para_xml(tmpl_name)
    base_rpr = _rpr(tmpl_name)
    _add_inline_runs(p_xml, text, base_rpr, doc_part,
                     force_bold=force_bold, force_italic=force_italic)
    doc.element.body.append(p_xml)
    return p_xml


# ── Document clearing ─────────────────────────────────────────────────────────

def _clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)


# ── Blockquote group parser ───────────────────────────────────────────────────

def _is_all_bold(text):
    return text.startswith('**') and text.endswith('**') and text.count('**') == 2


def _parse_blockquote_group(doc, bq_lines):
    """Split bq_lines on blank lines and insert one paragraph per group."""
    groups = [[]]
    for ln in bq_lines:
        if ln.strip() == '':
            groups.append([])
        else:
            groups[-1].append(ln)

    for group in groups:
        if not group:
            continue
        text = ' '.join(group)
        if _is_all_bold(text):
            _insert(doc, 'company', text[2:-2], force_bold=True)
        else:
            _insert(doc, 'role', text)


# ── Main line parser ──────────────────────────────────────────────────────────

_SPECIAL = ('# ', '- ', '> ', '***', '**')


def _parse_resume_lines(doc, lines):
    first = True
    contact_next = False

    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        # Section heading
        if line.startswith('# '):
            _insert(doc, 'heading', line[2:])
            i += 1
            continue

        # Bullet (collect indented continuation lines)
        if line.startswith('- '):
            text = line[2:]
            while i + 1 < len(lines) and lines[i + 1].startswith('  '):
                i += 1
                text += ' ' + lines[i].strip()
            _insert(doc, 'bullet', text)
            i += 1
            continue

        # Blockquote block
        if line.startswith('> ') or line == '>':
            bq_raw = []
            while i < len(lines) and (lines[i].startswith('> ') or lines[i] == '>'):
                bq_raw.append(lines[i][2:] if lines[i].startswith('> ') else '')
                i += 1
            if contact_next:
                contact_next = False
                text = ' '.join(ln for ln in bq_raw if ln.strip())
                _insert(doc, 'contact', text)
            else:
                _parse_blockquote_group(doc, bq_raw)
            continue

        # Subsection header ***text*** — use 'subsection' style if available.
        # Add 2 leading spaces when the template subsection has no left pPr
        # indent (comp template uses right-only indent; exp template uses left).
        if line.startswith('***') and line.endswith('***') and len(line) > 6:
            text = line[3:-3]
            sub_para = _tpara('subsection')
            if sub_para is not None:
                sub_ppr = sub_para._p.find(qn('w:pPr'))
                ind_el = sub_ppr.find(qn('w:ind')) if sub_ppr is not None else None
                has_left_indent = (
                    ind_el is not None
                    and ind_el.get(qn('w:left')) is not None
                )
                if not has_left_indent:
                    text = '  ' + text
                _insert(doc, 'subsection', text, force_bold=True, force_italic=True)
            else:
                _insert(doc, 'role', text, force_bold=True, force_italic=True)
            i += 1
            continue

        # Name (very first non-blank content line)
        if first:
            first = False
            contact_next = True
            text = line[2:-2] if (line.startswith('**') and line.endswith('**')) else line
            _insert(doc, 'name', text, force_bold=True)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            i += 1
            continue

        # Bold-lead line (conference header / org title)
        if line.startswith('**'):
            _insert(doc, 'conf', line)
            i += 1
            continue

        # Plain Body Text paragraph (collect wrapped continuation)
        text = line
        while i + 1 < len(lines):
            nxt = lines[i + 1]
            if not nxt.strip():
                break
            if nxt == '>' or any(nxt.startswith(s) for s in _SPECIAL):
                break
            i += 1
            text += ' ' + nxt.strip()
        _insert(doc, 'body', text)
        i += 1


# ── Public builders ───────────────────────────────────────────────────────────

def build_resume_docx_from_md(md_text, output_path, template_path=None):
    global _active_template
    _active_template = str(template_path or TEMPLATE)
    doc = Document(_active_template)
    _clear_body(doc)
    _parse_resume_lines(doc, md_text.splitlines())
    doc.save(str(output_path))
    _active_template = None


def build_resume_docx(summary_bullets, skills_md, sections, output_path,
                      template_path=None):
    import sys
    sys.path.insert(0, str(ROOT / 'scripts'))
    from draft_all import build_resume
    build_resume_docx_from_md(
        build_resume(summary_bullets, skills_md, sections),
        output_path,
        template_path=template_path,
    )


def build_cover_letter_docx(cover_letter_text, output_path):
    global _active_template
    _active_template = str(TEMPLATE)
    doc = Document(str(TEMPLATE))
    _clear_body(doc)
    for line in cover_letter_text.splitlines():
        if line.strip() == '---':
            continue
        _insert(doc, 'body', line)
    doc.save(str(output_path))
    _active_template = None
